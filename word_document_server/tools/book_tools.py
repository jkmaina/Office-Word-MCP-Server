"""
Book and front-matter tools for Word Document Server.

These tools generate title pages, copyright/legal pages, and
other front-matter sections.
"""
import os
from typing import Optional

from docx import Document

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def add_title_page(
    filename: str,
    title: str,
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
    date: Optional[str] = None,
) -> str:
    """Create a title page as a standalone new document or overwrite.

    Args:
        filename: Path to the new Word document (.docx)
        title: Main title text
        subtitle: Optional subtitle text
        author: Optional author name
        date: Optional date string
    """
    filename = ensure_docx_extension(filename)
    # Check if we can write/overwrite
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot write title page: {err}"
    try:
        doc = Document()
        # Title
        title_par = doc.add_paragraph()
        title_par.alignment = 1  # center
        run = title_par.add_run(title)
        run.bold = True
        run.font.size = run.font.size  # keep default large size
        # Subtitle
        if subtitle:
            sub_par = doc.add_paragraph()
            sub_par.alignment = 1
            sub_run = sub_par.add_run(subtitle)
            sub_run.italic = True
        # Author
        if author:
            auth_par = doc.add_paragraph()
            auth_par.alignment = 1
            auth_par.add_run(author)
        # Date
        if date:
            date_par = doc.add_paragraph()
            date_par.alignment = 1
            date_par.add_run(date)
        doc.save(filename)
        return f"Title page created: {filename}"
    except Exception as e:
        return f"Failed to create title page: {e}"


async def add_copyright_page(
    filename: str,
    text: str,
) -> str:
    """Append a copyright page at the end of a document.

    Args:
        filename: Path to an existing Word document
        text: Copyright or legal information text
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        doc.add_page_break()
        par = doc.add_paragraph(text)
        par.alignment = 1  # center
        doc.save(filename)
        return f"Copyright page appended to {filename}"
    except Exception as e:
        return f"Failed to append copyright page: {e}"


async def add_front_matter(
    filename: str,
    section: str,
    text: str,
) -> str:
    """Insert a front-matter section with a heading and text.

    Args:
        filename: Path to an existing Word document
        section: Name of the section (e.g., 'Preface')
        text: Paragraph text for the section
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        doc.add_page_break()
        doc.add_heading(section, level=1)
        doc.add_paragraph(text)
        doc.save(filename)
        return f"Front matter '{section}' added to {filename}"
    except Exception as e:
        return f"Failed to add front matter: {e}"