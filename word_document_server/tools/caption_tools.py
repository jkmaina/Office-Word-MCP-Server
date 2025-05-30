"""
Caption and list-generation tools for Word Document Server.

These tools insert numbered captions for figures/tables and generate
List of Figures/Tables based on those captions.
"""
import os
import json

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def insert_caption(
    filename: str,
    object_type: str,
    text: str,
    style: str = 'Caption',
) -> str:
    """Insert a caption for a figure or table and auto-number it.

    Args:
        filename: Path to an existing Word document
        object_type: 'Figure' or 'Table'
        text: Caption text (description)
        style: Style name to apply to the caption
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writable, err = check_file_writeable(filename)
    if not writable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        # Count existing captions of this type
        count = 0
        for p in doc.paragraphs:
            if p.text.startswith(object_type):
                count += 1
        number = count + 1
        caption_text = f"{object_type} {number}: {text}"
        par = doc.add_paragraph(caption_text)
        try:
            par.style = style
        except Exception:
            pass
        doc.save(filename)
        return f"Caption added: {caption_text}"
    except Exception as e:
        return f"Failed to insert caption: {e}"

async def generate_list_of_figures(
    filename: str,
    heading_style: str = 'TOC Heading',
) -> str:
    """Generate a List of Figures placeholder.

    Args:
        filename: Path to an existing Word document
        heading_style: Style name for the list heading
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writable, err = check_file_writeable(filename)
    if not writable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        doc.add_page_break()
        h = doc.add_paragraph('List of Figures')
        try:
            h.style = heading_style
        except Exception:
            pass
        p = doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'TOC \h \z \c "Figure"')
        p._p.append(fld)
        doc.save(filename)
        return f"List of Figures placeholder inserted into {filename}"
    except Exception as e:
        return f"Failed to insert List of Figures: {e}"

async def generate_list_of_tables(
    filename: str,
    heading_style: str = 'TOC Heading',
) -> str:
    """Generate a List of Tables placeholder.

    Args:
        filename: Path to an existing Word document
        heading_style: Style name for the list heading
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writable, err = check_file_writeable(filename)
    if not writable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        doc.add_page_break()
        h = doc.add_paragraph('List of Tables')
        try:
            h.style = heading_style
        except Exception:
            pass
        p = doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        fld.set(qn('w:instr'), 'TOC \h \z \c "Table"')
        p._p.append(fld)
        doc.save(filename)
        return f"List of Tables placeholder inserted into {filename}"
    except Exception as e:
        return f"Failed to insert List of Tables: {e}"