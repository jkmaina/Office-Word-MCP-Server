"""
Chapter tools for Word Document Server.

These tools manage automated chapter numbering and insertion.
"""
import os

from docx import Document
from docx.enum.section import WD_SECTION_START

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def new_chapter(
    filename: str,
    title: str,
    style: str = 'Heading 1',
) -> str:
    """Insert a new chapter with automatic numbering.

    The chapter number is determined by counting existing headings of the given style.

    Args:
        filename: Path to the Word document
        title: Chapter title text
        style: Style name for chapter headings (e.g., 'Heading 1')
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        # Determine next chapter number by counting existing headings
        count = 0
        for para in doc.paragraphs:
            try:
                if para.style.name == style and para.text.lower().startswith('chapter'):
                    count += 1
            except Exception:
                continue
        chapter_num = count + 1
        # Insert section break for new chapter
        try:
            doc.add_section(WD_SECTION_START.NEXT_PAGE)
        except Exception:
            doc.add_page_break()
        # Add chapter heading
        heading_text = f"Chapter {chapter_num}: {title}"
        try:
            doc.add_heading(heading_text, level=int(style.split()[-1]))
        except Exception:
            para = doc.add_paragraph(heading_text)
            para.style = style
        doc.save(filename)
        return f"{heading_text} added to {filename}"
    except Exception as e:
        return f"Failed to add new chapter: {e}"