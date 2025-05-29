"""
Code block insertion tools for Word Document Server.

These tools insert formatted code listings (monospaced font with shading)
as tables into Word documents.
"""
import os
from typing import Optional

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def add_code_block(
    filename: str,
    code_text: str,
    language: Optional[str] = None,
    style: str = 'CodeBlock',
) -> str:
    """Insert a code block into a Word document.

    Args:
        filename: Path to the Word document
        code_text: The code content (can include '\n' for multiple lines)
        language: Optional language name (ignored, placeholder)
        style: Paragraph style name to apply (will be created if missing)
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        # Ensure paragraph style exists
        styles = doc.styles
        if style not in styles:
            # Create a new paragraph style based on Normal
            new_style = styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)
            font = new_style.font
            font.name = 'Courier New'
            font.size = None
        # Insert a one-cell table for shading
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        # Apply shading to cell
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
        # Populate code text
        cell.text = code_text
        # Apply style and monospaced font to each paragraph
        for paragraph in cell.paragraphs:
            try:
                paragraph.style = style
            except Exception:
                pass
            for run in paragraph.runs:
                run.font.name = 'Courier New'
        doc.save(filename)
        return f"Code block added to {filename}"
    except Exception as e:
        return f"Failed to add code block: {e}"