"""
Header and footer tools for Word Document Server.

These tools insert headers and footers with optional dynamic fields
(e.g. page numbers) into Word documents.
"""
import os
from typing import Optional, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def insert_header(
    filename: str,
    text: str,
    alignment: str = 'center',
    fields: Optional[List[str]] = None,
) -> str:
    """Insert a header into all sections of a Word document.

    Args:
        filename: Path to the Word document
        text: Header text containing optional placeholders {pagenum}
        alignment: 'left', 'center', or 'right'
        fields: List of dynamic fields to insert, e.g. ['pagenum']
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        for section in doc.sections:
            hdr = section.header
            # Break linkage to previous header
            try:
                hdr.is_linked_to_previous = False
            except Exception:
                pass
            # Add a paragraph for the header
            para = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
            para.clear()
            # Set alignment
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            # Add text and fields
            parts = text.split('{pagenum}') if fields and 'pagenum' in fields else [text]
            for i, part in enumerate(parts):
                if part:
                    para.add_run(part)
                if fields and 'pagenum' in fields and i < len(parts) - 1:
                    fld = OxmlElement('w:fldSimple')
                    fld.set(qn('w:instr'), 'PAGE')
                    para._p.append(fld)
        doc.save(filename)
        return f"Header inserted into {filename}"
    except Exception as e:
        return f"Failed to insert header: {e}"


async def insert_footer(
    filename: str,
    text: str,
    alignment: str = 'center',
    fields: Optional[List[str]] = None,
) -> str:
    """Insert a footer into all sections of a Word document.

    Args:
        filename: Path to the Word document
        text: Footer text containing optional placeholders {pagenum}
        alignment: 'left', 'center', or 'right'
        fields: List of dynamic fields to insert, e.g. ['pagenum']
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        for section in doc.sections:
            ftr = section.footer
            try:
                ftr.is_linked_to_previous = False
            except Exception:
                pass
            para = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
            para.clear()
            align_map = {
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
            }
            para.alignment = align_map.get(alignment.lower(), WD_ALIGN_PARAGRAPH.CENTER)
            parts = text.split('{pagenum}') if fields and 'pagenum' in fields else [text]
            for i, part in enumerate(parts):
                if part:
                    para.add_run(part)
                if fields and 'pagenum' in fields and i < len(parts) - 1:
                    fld = OxmlElement('w:fldSimple')
                    fld.set(qn('w:instr'), 'PAGE')
                    para._p.append(fld)
        doc.save(filename)
        return f"Footer inserted into {filename}"
    except Exception as e:
        return f"Failed to insert footer: {e}"