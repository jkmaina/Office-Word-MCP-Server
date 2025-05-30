"""
Navigation tools for Word Document Server.

These tools insert table of contents placeholders and manage cross-document navigation.
"""
import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

from word_document_server.utils.file_utils import ensure_docx_extension, check_file_writeable


async def insert_toc_placeholder(
    filename: str,
    heading_style: str = 'TOC Heading',
) -> str:
    """Insert a Table of Contents field placeholder.

    Args:
        filename: Path to an existing Word document
        heading_style: Style name to apply to the TOC heading
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writeable, err = check_file_writeable(filename)
    if not writeable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        # Insert heading for TOC
        toc_heading = doc.add_paragraph('Table of Contents')
        try:
            toc_heading.style = heading_style
        except Exception:
            pass
        # Insert the TOC field
        p = doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        # Classic TOC field code (levels 1-3, hyperlinks, hide page numbers)
        fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
        p._p.append(fld)
        doc.save(filename)
        return f"TOC placeholder inserted into {filename}"
    except Exception as e:
        return f"Failed to insert TOC placeholder: {e}"
    
async def bookmark(
    filename: str,
    bookmark_name: str,
) -> str:
    """Insert a bookmark at the start of the document for cross-referencing.

    Args:
        filename: Path to an existing Word document
        bookmark_name: Name of the bookmark (must start with a letter)
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writable, err = check_file_writeable(filename)
    if not writable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        p = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        # Create bookmark start
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_name)
        # Create bookmark end
        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        # Insert elements
        p._p.insert(0, start)
        p._p.insert(1, end)
        doc.save(filename)
        return f"Bookmark '{bookmark_name}' added to {filename}"
    except Exception as e:
        return f"Failed to add bookmark: {e}"

async def insert_hyperlink(
    filename: str,
    bookmark_name: str,
    display_text: str,
) -> str:
    """Insert a hyperlink to a bookmark in the document.

    Args:
        filename: Path to an existing Word document
        bookmark_name: Name of the target bookmark
        display_text: Text to display for the hyperlink
    """
    filename = ensure_docx_extension(filename)
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    writable, err = check_file_writeable(filename)
    if not writable:
        return f"Cannot modify document: {err}"
    try:
        doc = Document(filename)
        # Append hyperlink at end of document
        p = doc.add_paragraph()
        fld = OxmlElement('w:fldSimple')
        # Internal link to bookmark
        fld.set(qn('w:instr'), rf'HYPERLINK \l "{bookmark_name}"')
        # Add display text child node
        run = OxmlElement('w:r')
        text = OxmlElement('w:t')
        text.text = display_text
        run.append(text)
        fld.append(run)
        p._p.append(fld)
        doc.save(filename)
        return f"Hyperlink to '{bookmark_name}' added as '{display_text}'"
    except Exception as e:
        return f"Failed to insert hyperlink: {e}"